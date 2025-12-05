"""OCR service for extracting text from images and PDFs."""

import io
import os
import tempfile
from abc import ABC, abstractmethod
from dataclasses import dataclass
from typing import List, Optional, Tuple
import structlog

from PIL import Image
import pytesseract
from pdf2image import convert_from_bytes
import pdfplumber
from docx import Document

from ..core.config import get_settings

logger = structlog.get_logger()


@dataclass
class OCRResult:
    """Result of OCR processing."""
    text: str
    confidence: float
    provider: str
    pages_processed: int = 1
    raw_data: Optional[dict] = None


class OCRService(ABC):
    """Abstract base class for OCR services."""
    
    @abstractmethod
    def extract_text(self, content: bytes, content_type: str, filename: str) -> OCRResult:
        """Extract text from file content."""
        pass


class TesseractOCR(OCRService):
    """Tesseract-based OCR implementation."""
    
    def __init__(self):
        self.settings = get_settings()
        
        # Set tesseract command if specified
        if self.settings.tesseract_cmd:
            pytesseract.pytesseract.tesseract_cmd = self.settings.tesseract_cmd
    
    def extract_text(self, content: bytes, content_type: str, filename: str) -> OCRResult:
        """Extract text from file content using Tesseract."""
        try:
            if content_type == 'application/pdf' or filename.lower().endswith('.pdf'):
                return self._process_pdf(content)
            elif content_type.startswith('image/') or self._is_image_file(filename):
                return self._process_image(content)
            elif 'spreadsheet' in content_type or filename.lower().endswith(('.xlsx', '.xls')):
                return self._process_excel(content, filename)
            elif content_type == 'text/csv' or filename.lower().endswith('.csv'):
                return self._process_csv(content)
            elif 'wordprocessing' in content_type or filename.lower().endswith(('.docx', '.doc')):
                return self._process_docx(content)
            else:
                logger.warning(
                    "Unsupported content type for OCR",
                    content_type=content_type,
                    filename=filename
                )
                return OCRResult(
                    text="",
                    confidence=0.0,
                    provider="tesseract"
                )
        except Exception as e:
            logger.error("OCR extraction failed", error=str(e), filename=filename)
            return OCRResult(
                text="",
                confidence=0.0,
                provider="tesseract"
            )
    
    def _process_pdf(self, content: bytes) -> OCRResult:
        """Process PDF file."""
        all_text = []
        confidences = []
        pages_processed = 0
        
        # First, try to extract text directly (for digital PDFs)
        try:
            with pdfplumber.open(io.BytesIO(content)) as pdf:
                for page in pdf.pages:
                    text = page.extract_text()
                    if text and text.strip():
                        all_text.append(text)
                        confidences.append(0.95)  # High confidence for extracted text
                        pages_processed += 1
        except Exception as e:
            logger.debug("Direct PDF text extraction failed, trying OCR", error=str(e))
        
        # If no text extracted, use OCR
        if not all_text:
            try:
                # Convert PDF to images
                images = convert_from_bytes(content, dpi=300)
                
                for i, image in enumerate(images):
                    try:
                        # Run OCR with confidence data
                        data = pytesseract.image_to_data(
                            image, 
                            output_type=pytesseract.Output.DICT
                        )
                        
                        # Extract text and confidence
                        page_text = []
                        page_confidences = []
                        
                        for j, text in enumerate(data['text']):
                            if text.strip():
                                page_text.append(text)
                                conf = data['conf'][j]
                                if conf > 0:  # -1 means no confidence
                                    page_confidences.append(conf / 100.0)
                        
                        if page_text:
                            all_text.append(' '.join(page_text))
                            if page_confidences:
                                confidences.append(sum(page_confidences) / len(page_confidences))
                        
                        pages_processed += 1
                        
                    except Exception as e:
                        logger.error(f"OCR failed for page {i}", error=str(e))
                        
            except Exception as e:
                logger.error("PDF to image conversion failed", error=str(e))
        
        combined_text = '\n\n'.join(all_text)
        avg_confidence = sum(confidences) / len(confidences) if confidences else 0.0
        
        return OCRResult(
            text=combined_text,
            confidence=avg_confidence,
            provider="tesseract",
            pages_processed=pages_processed
        )
    
    def _process_image(self, content: bytes) -> OCRResult:
        """Process image file with OCR."""
        try:
            image = Image.open(io.BytesIO(content))
            
            # Preprocess for better OCR
            if image.mode != 'RGB':
                image = image.convert('RGB')
            
            # Get OCR data with confidence
            data = pytesseract.image_to_data(
                image,
                output_type=pytesseract.Output.DICT
            )
            
            # Extract text and confidence
            texts = []
            confidences = []
            
            for i, text in enumerate(data['text']):
                if text.strip():
                    texts.append(text)
                    conf = data['conf'][i]
                    if conf > 0:
                        confidences.append(conf / 100.0)
            
            combined_text = ' '.join(texts)
            avg_confidence = sum(confidences) / len(confidences) if confidences else 0.0
            
            return OCRResult(
                text=combined_text,
                confidence=avg_confidence,
                provider="tesseract",
                pages_processed=1
            )
            
        except Exception as e:
            logger.error("Image OCR failed", error=str(e))
            return OCRResult(text="", confidence=0.0, provider="tesseract")
    
    def _process_excel(self, content: bytes, filename: str) -> OCRResult:
        """Extract text from Excel files."""
        try:
            import pandas as pd
            
            # Read all sheets
            xlsx = pd.ExcelFile(io.BytesIO(content))
            all_text = []
            
            for sheet_name in xlsx.sheet_names:
                df = pd.read_excel(xlsx, sheet_name=sheet_name)
                
                # Convert to text representation
                text_parts = [f"Sheet: {sheet_name}"]
                
                # Add headers
                headers = [str(col) for col in df.columns]
                text_parts.append("Headers: " + ", ".join(headers))
                
                # Add data rows
                for _, row in df.iterrows():
                    row_text = " | ".join([str(v) for v in row.values if pd.notna(v)])
                    if row_text.strip():
                        text_parts.append(row_text)
                
                all_text.append('\n'.join(text_parts))
            
            return OCRResult(
                text='\n\n'.join(all_text),
                confidence=0.98,  # High confidence for structured data
                provider="pandas",
                pages_processed=len(xlsx.sheet_names)
            )
            
        except Exception as e:
            logger.error("Excel extraction failed", error=str(e))
            return OCRResult(text="", confidence=0.0, provider="pandas")
    
    def _process_csv(self, content: bytes) -> OCRResult:
        """Extract text from CSV files."""
        try:
            import pandas as pd
            
            # Try different encodings
            text_content = None
            for encoding in ['utf-8', 'latin-1', 'cp1252']:
                try:
                    text_content = content.decode(encoding)
                    break
                except UnicodeDecodeError:
                    continue
            
            if not text_content:
                return OCRResult(text="", confidence=0.0, provider="pandas")
            
            df = pd.read_csv(io.StringIO(text_content))
            
            # Convert to text
            text_parts = []
            
            # Headers
            headers = [str(col) for col in df.columns]
            text_parts.append("Headers: " + ", ".join(headers))
            
            # Data rows
            for _, row in df.iterrows():
                row_text = " | ".join([str(v) for v in row.values if pd.notna(v)])
                if row_text.strip():
                    text_parts.append(row_text)
            
            return OCRResult(
                text='\n'.join(text_parts),
                confidence=0.98,
                provider="pandas",
                pages_processed=1
            )
            
        except Exception as e:
            logger.error("CSV extraction failed", error=str(e))
            return OCRResult(text="", confidence=0.0, provider="pandas")
    
    def _process_docx(self, content: bytes) -> OCRResult:
        """Extract text from DOCX files."""
        try:
            doc = Document(io.BytesIO(content))
            
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            
            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text for cell in row.cells if cell.text.strip()])
                    if row_text.strip():
                        paragraphs.append(row_text)
            
            return OCRResult(
                text='\n'.join(paragraphs),
                confidence=0.98,
                provider="python-docx",
                pages_processed=1
            )
            
        except Exception as e:
            logger.error("DOCX extraction failed", error=str(e))
            return OCRResult(text="", confidence=0.0, provider="python-docx")
    
    def _is_image_file(self, filename: str) -> bool:
        """Check if filename indicates an image file."""
        image_extensions = {'.jpg', '.jpeg', '.png', '.gif', '.bmp', '.tiff', '.tif'}
        ext = os.path.splitext(filename)[1].lower()
        return ext in image_extensions


def get_ocr_service() -> OCRService:
    """Factory function to get the OCR service."""
    settings = get_settings()
    
    # Currently only Tesseract is implemented
    # Could add Google Vision or AWS Textract here
    return TesseractOCR()

