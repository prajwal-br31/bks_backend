import csv
import io
import logging
import re
from dataclasses import dataclass, field
from typing import Optional

logger = logging.getLogger(__name__)


@dataclass
class ExtractedContent:
    """Represents extracted content from a document."""
    text: str
    tables: list[list[list[str]]] = field(default_factory=list)  # List of tables, each table is rows x cols
    metadata: dict = field(default_factory=dict)
    page_count: int = 0
    confidence: float = 1.0  # OCR confidence if applicable


class ContentExtractor:
    """
    Extracts text content from various document types.
    
    Handles:
    - PDF (via PyPDF2 or pdf2image + OCR)
    - XLSX/XLS (via openpyxl/xlrd)
    - CSV
    - DOCX (via python-docx)
    - Images (delegated to OCR service)
    """

    def extract_pdf(self, content: bytes) -> ExtractedContent:
        """Extract text from PDF files."""
        try:
            import pypdf
            
            reader = pypdf.PdfReader(io.BytesIO(content))
            text_parts = []
            
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)
            
            full_text = "\n\n".join(text_parts)
            
            # If no text extracted, it might be a scanned PDF
            if not full_text.strip():
                return ExtractedContent(
                    text="",
                    metadata={"needs_ocr": True, "page_count": len(reader.pages)},
                    page_count=len(reader.pages),
                    confidence=0.0,
                )
            
            return ExtractedContent(
                text=full_text,
                page_count=len(reader.pages),
                metadata={"pdf_info": reader.metadata._data if reader.metadata else {}},
            )
        except ImportError:
            logger.error("pypdf not installed. Run: pip install pypdf")
            return ExtractedContent(text="", metadata={"error": "pypdf not installed"})
        except Exception as e:
            logger.error(f"Error extracting PDF: {e}")
            return ExtractedContent(text="", metadata={"error": str(e)})

    def extract_xlsx(self, content: bytes) -> ExtractedContent:
        """Extract text and tables from Excel files."""
        try:
            import openpyxl
            
            workbook = openpyxl.load_workbook(io.BytesIO(content), data_only=True)
            text_parts = []
            tables = []
            
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                text_parts.append(f"Sheet: {sheet_name}")
                
                table = []
                for row in sheet.iter_rows(values_only=True):
                    # Convert all values to strings
                    str_row = [str(cell) if cell is not None else "" for cell in row]
                    if any(str_row):  # Skip completely empty rows
                        table.append(str_row)
                        text_parts.append(" | ".join(str_row))
                
                if table:
                    tables.append(table)
            
            return ExtractedContent(
                text="\n".join(text_parts),
                tables=tables,
                metadata={"sheet_count": len(workbook.sheetnames), "sheets": workbook.sheetnames},
            )
        except ImportError:
            logger.error("openpyxl not installed. Run: pip install openpyxl")
            return ExtractedContent(text="", metadata={"error": "openpyxl not installed"})
        except Exception as e:
            logger.error(f"Error extracting XLSX: {e}")
            return ExtractedContent(text="", metadata={"error": str(e)})

    def extract_csv(self, content: bytes) -> ExtractedContent:
        """Extract text and tables from CSV files."""
        try:
            # Try to detect encoding
            text_content = self._decode_content(content)
            
            reader = csv.reader(io.StringIO(text_content))
            table = list(reader)
            
            # Build text representation
            text_parts = []
            for row in table:
                text_parts.append(" | ".join(row))
            
            return ExtractedContent(
                text="\n".join(text_parts),
                tables=[table] if table else [],
                metadata={"row_count": len(table)},
            )
        except Exception as e:
            logger.error(f"Error extracting CSV: {e}")
            return ExtractedContent(text="", metadata={"error": str(e)})

    def extract_docx(self, content: bytes) -> ExtractedContent:
        """Extract text from Word documents."""
        try:
            from docx import Document
            
            doc = Document(io.BytesIO(content))
            text_parts = []
            tables = []
            
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)
                    text_parts.append(" | ".join(row_data))
                tables.append(table_data)
            
            return ExtractedContent(
                text="\n".join(text_parts),
                tables=tables,
                metadata={"paragraph_count": len(doc.paragraphs), "table_count": len(doc.tables)},
            )
        except ImportError:
            logger.error("python-docx not installed. Run: pip install python-docx")
            return ExtractedContent(text="", metadata={"error": "python-docx not installed"})
        except Exception as e:
            logger.error(f"Error extracting DOCX: {e}")
            return ExtractedContent(text="", metadata={"error": str(e)})

    def extract_from_html(self, html: str) -> ExtractedContent:
        """Extract text and tables from HTML content (email body)."""
        try:
            from bs4 import BeautifulSoup
            
            soup = BeautifulSoup(html, "html.parser")
            
            # Remove script and style elements
            for element in soup(["script", "style", "head"]):
                element.decompose()
            
            # Extract tables
            tables = []
            for table in soup.find_all("table"):
                table_data = []
                for row in table.find_all("tr"):
                    cells = row.find_all(["td", "th"])
                    row_data = [cell.get_text(strip=True) for cell in cells]
                    if any(row_data):
                        table_data.append(row_data)
                if table_data:
                    tables.append(table_data)
            
            # Get text
            text = soup.get_text(separator="\n", strip=True)
            
            # Extract links to downloadable files
            links = []
            for a in soup.find_all("a", href=True):
                href = a["href"]
                if any(ext in href.lower() for ext in [".pdf", ".xlsx", ".xls", ".csv", ".doc"]):
                    links.append({"url": href, "text": a.get_text(strip=True)})
            
            return ExtractedContent(
                text=text,
                tables=tables,
                metadata={"links": links, "has_tables": len(tables) > 0},
            )
        except ImportError:
            logger.error("beautifulsoup4 not installed. Run: pip install beautifulsoup4")
            return ExtractedContent(text=html, metadata={"error": "beautifulsoup4 not installed"})
        except Exception as e:
            logger.error(f"Error extracting HTML: {e}")
            return ExtractedContent(text=html, metadata={"error": str(e)})

    def extract(self, content: bytes, content_type: str, filename: str = "") -> ExtractedContent:
        """
        Extract content based on file type.
        
        Args:
            content: Binary file content
            content_type: MIME type
            filename: Original filename (for extension-based detection)
        
        Returns:
            ExtractedContent with text and metadata
        """
        content_type = content_type.lower()
        filename_lower = filename.lower()

        # PDF
        if content_type == "application/pdf" or filename_lower.endswith(".pdf"):
            return self.extract_pdf(content)

        # Excel
        if "spreadsheet" in content_type or filename_lower.endswith((".xlsx", ".xls")):
            return self.extract_xlsx(content)

        # CSV
        if content_type == "text/csv" or filename_lower.endswith(".csv"):
            return self.extract_csv(content)

        # Word
        if "wordprocessing" in content_type or filename_lower.endswith((".docx", ".doc")):
            return self.extract_docx(content)

        # Images - return empty, needs OCR
        if content_type.startswith("image/") or filename_lower.endswith((".jpg", ".jpeg", ".png", ".gif", ".tiff")):
            return ExtractedContent(
                text="",
                metadata={"needs_ocr": True, "content_type": content_type},
                confidence=0.0,
            )

        # Unknown type
        return ExtractedContent(
            text="",
            metadata={"unsupported_type": content_type},
        )

    def _decode_content(self, content: bytes) -> str:
        """Try to decode bytes content with various encodings."""
        encodings = ["utf-8", "latin-1", "cp1252", "iso-8859-1"]
        
        for encoding in encodings:
            try:
                return content.decode(encoding)
            except UnicodeDecodeError:
                continue
        
        # Last resort: decode with errors replaced
        return content.decode("utf-8", errors="replace")

